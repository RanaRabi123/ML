import os 
import langchain_core
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage
from langchain_chroma import Chroma
from langchain_core.documents import Document as LCDocument
from langchain_core.runnables import RunnablePassthrough, RunnableLambda
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from pypdf import PdfReader
from dotenv import load_dotenv
from docx import Document

load_dotenv()
model = ChatGroq(
    model='llama-3.3-70b-versatile'
)



def preprocess():
    path = r'D:\intern preparation\Langchain\study mate.docx'
    if path.endswith('.docx'):
        doc = Document(path)
        full_text = '\n'.join([p.text for p in doc.paragraphs])
    elif path.endswith('.pdf'):
        file = PdfReader(path)
        full_text = '\n\n'.join([i.extract_text() or '' for i in file.pages])
    else:
        raise TypeError('Unexpected file type... ')

    return full_text

def split():
    full_text = preprocess()
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500, 
        chunk_overlap=50
    )

    document = [LCDocument(page_content=full_text)]
    chunks = text_splitter.split_documents(document)

    print('total chunk are : ', len(chunks))
    return chunks


def embed():
    embed_model = HuggingFaceEmbeddings(model_name='neuml/pubmedbert-base-embeddings')
    persist_dir = r'D:\intern preparation\Langchain\chroma_db'  # use a dedicated subfolder

    if os.path.exists(persist_dir) and os.listdir(persist_dir):
        local_db = Chroma(persist_directory=persist_dir, embedding_function=embed_model)
    else:
        local_db = Chroma.from_documents(
            documents=split(),
            embedding=embed_model,
            persist_directory=persist_dir
        )
    return local_db

local_db = embed()

def retrieve():
    retrieval = local_db.as_retriever(search_kwargs={'k':2})
    return retrieval

def format_doc(docs):
    return '\n\n'.join(doc.page_content for doc in docs)

chat_memory=[]
max_len_memory=15

prompt = ChatPromptTemplate.from_messages([
        ('system', 'You are a helpful assistant. Use the following given info/context to give answer to user \n{context} .If you dont know anything just say i dont know, also use the chat history to assist it too'),
        MessagesPlaceholder(variable_name='chat_memory')
        ('user', '{question}')
    ])

chain = {'context': retrieve() | RunnableLambda(format_doc), 'chat_memory': RunnableLambda(lambda x: chat_memory), 'question':RunnablePassthrough()} | prompt|model|StrOutputParser()

def summarize_memory(to_summarized):
    summarize_prompt = ChatPromptTemplate.from_messages([
        ('system', 'You are a  chat summarizer, summarize it but keep key facts,idea and important things'),
        MessagesPlaceholder(variable_name='memory')
    ])

    return (summarize_prompt|model).invoke({'memory':to_summarized}).content



while True:
    user = input('ask anything : ')

    if user.lower().strip() in ['exit', 'quit']:
        break
    else:
        if len(chat_memory)>max_len_memory:
            to_keep=chat_memory[-2:]
            to_summarized= chat_memory[:-2]
            summary_chat= summarize_memory(to_summarized)
            chat_memory = [SystemMessage(content=f'summary of earlier conversations {summary_chat}')]+to_keep
        
        try:
            chat_memory.append(HumanMessage(content=user))
            response = chain.invoke(
                    user
                )
                
            chat_memory.append(AIMessage(content=response))
            print('Ai : ', response)
        except Exception as e:
            print(f'An error occured while processing {e}')

